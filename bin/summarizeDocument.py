import os
import re
import argparse
import textwrap
import json
import requests
import time
from pypdf import PdfReader
from docx import Document
from typing import List, Dict, Any

# A default word limit to trigger chunking.
# This is a safe estimate for the llama3 8B model's ~8k token context window.
DEFAULT_WORD_LIMIT = 5000
CHUNK_SIZE = 2000

def check_ollama() -> bool:
    """Checks if the Ollama server is running and the llama3 model is available."""
    try:
        response = requests.get("http://localhost:11434/api/tags")
        if response.status_code == 200:
            models = response.json().get("models", [])
            if any(model["name"].startswith("llama3") for model in models):
                print("Ollama server running with llama3 model available.")
                return True
            else:
                print("Ollama server running, but 'llama3' model not found.")
                print("Please run 'ollama pull llama3' to download it.")
                return False
        else:
            print(f"Ollama server returned status {response.status_code}.")
            return False
    except requests.ConnectionError:
        print("Error: Ollama server not running. Please start it with 'ollama serve'.")
        return False

def extract_text_and_metadata(file_path: str) -> (str, int, int):
    """Extracts text, page/paragraph count, and word count from a PDF, TXT, or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.pdf':
            reader = PdfReader(file_path)
            num_pages = len(reader.pages)
            text = ""
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            num_pages = 1
        elif ext == '.docx':
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            num_pages = len([p for p in doc.paragraphs if p.text.strip()])  # Approximate page count based on non-empty paragraphs
        else:
            print("Unsupported file type. Supported: .pdf, .txt, .docx")
            return None, 0, 0

        word_count = len(re.findall(r'\b\w+\b', text))
        return text, num_pages, word_count
    except Exception as e:
        print(f"Error extracting text and metadata: {e}")
        return None, 0, 0

def chunk_text(text: str, max_words: int = CHUNK_SIZE) -> List[str]:
    """Splits text into chunks of a maximum word count."""
    words = text.split()
    return [' '.join(words[i:i + max_words]) for i in range(0, len(words), max_words)]

def summarize_with_ollama(text: str, is_chunk: bool = False) -> (str, List[str]):
    """
    Sends text to Ollama for summarization. Handles large text by chunking.

    If is_chunk is True, it returns only the summary paragraph for a single chunk.
    Otherwise, it returns the final executive summary and key points.
    """
    word_count = len(re.findall(r'\b\w+\b', text))

    # Handle large documents by chunking and recursively summarizing
    if word_count > DEFAULT_WORD_LIMIT and not is_chunk:
        print(f"Document is large ({word_count} words). Splitting into chunks...", flush=True)
        chunks = chunk_text(text, max_words=CHUNK_SIZE)
        chunk_summaries = []

        for i, chunk in enumerate(chunks):
            print(f"Summarizing chunk {i+1}/{len(chunks)}...", flush=True)
            chunk_summary, _ = summarize_with_ollama(chunk, is_chunk=True)
            chunk_summaries.append(chunk_summary)

        combined_text = "\n\n".join(chunk_summaries)
        print("Generating final summary from chunk summaries...", flush=True)
        return summarize_with_ollama(combined_text)

    # Base case: summarize a single chunk or a small document
    if is_chunk:
        prompt = textwrap.dedent(f"""\
            Summarize the following document chunk in detail. The summary should be a comprehensive paragraph covering all key aspects.

            ---
            DOCUMENT CHUNK:
            {text}
            ---

            Please respond with only the summary paragraph.
            """)
    else:
        # Final summary prompt
        prompt = textwrap.dedent(f"""\
            You are a helpful assistant. Your task is to extract a detailed executive summary and a comprehensive list of key points from the provided document text.

            ---
            DOCUMENT TEXT:
            {text}
            ---

            Please respond with a detailed paragraph (200-400 words) for the executive summary, followed by a list of key points. Make the executive summary thorough and insightful, covering main themes, arguments, and conclusions. For key points, provide up to 30 detailed bullet points, each elaborating on important elements with context or explanations where relevant. Use the following format:

            Executive Summary:
            [Write the detailed summary paragraph here.]

            Key Points:
            - [List the first key point here, with detail.]
            - [List the second key point here, with detail.]
            - [Continue with up to 30 key points, using bullet points.]
            """)

    payload = {
        "model": "llama3",
        "prompt": prompt,
        "stream": False
    }

    try:
        response = requests.post("http://localhost:11434/api/generate", json=payload, timeout=300)
        response.raise_for_status() # Raise an exception for bad status codes

        data = response.json()
        raw_output = data["response"].strip()
        print(f"\n--- Raw Ollama Output for Parsing ---\n{raw_output}\n-----------------------------------\n", flush=True)

        # If summarizing a chunk, just return the raw output.
        if is_chunk:
            return raw_output, []

        # New, more robust parsing logic
        summary_start_marker = "Executive Summary:"
        key_points_start_marker = "Key Points:"

        summary_paragraph = ""
        key_points = []

        # Use regex to find and split the sections
        match = re.search(f"(?:^|\n\s*){re.escape(summary_start_marker)}(.*?)(?:^|\n\s*){re.escape(key_points_start_marker)}(.*)", raw_output, re.DOTALL | re.IGNORECASE)

        if match:
            # Extract and clean up the summary paragraph
            summary_paragraph = match.group(1).strip()
            summary_paragraph = re.sub(r'^\s*\*\*(.*?)\*\*\s*$', r'\1', summary_paragraph, flags=re.MULTILINE).strip()

            # Extract and parse the key points
            points_text = match.group(2).strip()
            points_list = re.findall(r'^\s*(?:-|\*|•|\d+\.)\s*(.*?)$', points_text, re.MULTILINE)
            key_points = [re.sub(r'^\s*\*+\s*(.*?)\s*\*+\s*$', r'\1', point).strip() for point in points_list if point.strip()]
        else:
            # Fallback for when key points are not found, but a summary is
            summary_match = re.search(f"(?:^|\n\s*){re.escape(summary_start_marker)}(.*)", raw_output, re.DOTALL | re.IGNORECASE)
            if summary_match:
                summary_paragraph = summary_match.group(1).strip()
                summary_paragraph = re.sub(r'^\s*\*\*(.*?)\*\*\s*$', r'\1', summary_paragraph, flags=re.MULTILINE).strip()


        return summary_paragraph, key_points

    except requests.exceptions.RequestException as e:
        print(f"Error communicating with Ollama API: {e}", flush=True)
        return "", []
    except Exception as e:
        print(f"An unexpected error occurred during summarization: {e}", flush=True)
        return "", []

def main():
    parser = argparse.ArgumentParser(description="Generate an executive summary from a PDF, TXT, or DOCX document using the local Llama3 model via a direct Ollama API call.")
    parser.add_argument("--file", type=str, help="Path to the PDF, TXT, or DOCX document")
    args = parser.parse_args()

    file_path = args.file
    # If the file path is not provided as a command-line argument, prompt the user for it.
    if file_path is None:
        file_path = input("Please enter the path to the PDF, TXT, or DOCX document: ")

    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    # Get the directory of the input file
    input_dir = os.path.dirname(os.path.abspath(file_path))

    if not check_ollama():
        print("\nAborting summarization.")
        return

    filename = os.path.basename(file_path)
    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
    text, num_pages, num_words = extract_text_and_metadata(file_path)
    if text is None or not text.strip():
        print("Failed to extract meaningful text from file. Aborting.")
        return

    print(f"\nDocument word count: {num_words}")
    print("\nRunning Ollama summarization...")
    start_time = time.time()
    summary_paragraph, key_points = summarize_with_ollama(text)
    end_time = time.time()
    time_taken = end_time - start_time

    print("\n--- Executive Summary ---")
    if summary_paragraph:
        print(summary_paragraph)
    else:
        print("No executive summary found.")

    print("\n--- Key Points ---")
    if key_points:
        for point in key_points:
            print(f"- {point}")
    else:
        print("No key points found.")

    base_name = os.path.splitext(filename)[0]
    output_filename = os.path.join(input_dir, f"{base_name}_executive_summary.txt")
    with open(output_filename, "w") as f:
        f.write("Executive Summary:\n")
        f.write(summary_paragraph)
        f.write("\n\nKey Points:\n")
        for point in key_points:
            f.write(f"- {point}\n")
    print(f"\nFormatted summary saved to {output_filename}")

    print(f"\nFilename: {filename}")
    print(f"File size: {file_size_mb:.2f} MB")
    print(f"Number of pages/sections: {num_pages}")
    print(f"Number of words: {num_words}")
    print(f"Time taken for Ollama API call: {time_taken:.2f} seconds")

if __name__ == "__main__":
    main()